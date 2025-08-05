import os
import streamlit as st
import PyPDF2
import pdfplumber
from docx import Document
import textract
from pathlib import Path
import config

class FileHandler:
    """Handle file uploads and text extraction"""
    
    def __init__(self):
        self.allowed_extensions = config.ALLOWED_EXTENSIONS
        self.max_file_size = config.MAX_FILE_SIZE_MB * 1024 * 1024  # Convert to bytes
    
    def validate_file(self, uploaded_file):
        """Validate uploaded file"""
        if uploaded_file is None:
            return False, "No file uploaded"
        
        # Check file size
        if uploaded_file.size > self.max_file_size:
            return False, f"File size exceeds {config.MAX_FILE_SIZE_MB}MB limit"
        
        # Check file extension
        file_extension = Path(uploaded_file.name).suffix.lower()
        if file_extension not in self.allowed_extensions:
            return False, f"File type {file_extension} not supported. Allowed types: {', '.join(self.allowed_extensions)}"
        
        return True, "File is valid"
    
    def extract_text_from_pdf(self, pdf_file, method='pdfplumber'):
        """Extract text from PDF file using different methods"""
        text = ""
        
        try:
            if method == 'pdfplumber':
                # Using pdfplumber (better for complex layouts)
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            
            elif method == 'pypdf2':
                # Using PyPDF2 (fallback method)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
        except Exception as e:
            st.error(f"Error extracting text from PDF: {str(e)}")
            # Try alternative method
            if method == 'pdfplumber':
                return self.extract_text_from_pdf(pdf_file, method='pypdf2')
        
        return text.strip()
    
    def extract_text_from_docx(self, docx_file):
        """Extract text from DOCX file"""
        text = ""
        
        try:
            doc = Document(docx_file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                        
        except Exception as e:
            st.error(f"Error extracting text from DOCX: {str(e)}")
        
        return text.strip()
    
    def extract_text_from_doc(self, doc_file):
        """Extract text from DOC file using textract"""
        text = ""
        
        try:
            # Save uploaded file temporarily
            temp_file_path = f"temp_{doc_file.name}"
            with open(temp_file_path, "wb") as f:
                f.write(doc_file.getbuffer())
            
            # Extract text using textract
            text = textract.process(temp_file_path).decode('utf-8')
            
            # Clean up temporary file
            os.remove(temp_file_path)
            
        except Exception as e:
            st.error(f"Error extracting text from DOC: {str(e)}")
            # Clean up temporary file if it exists
            temp_file_path = f"temp_{doc_file.name}"
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
        
        return text.strip()
    
    def extract_text_from_txt(self, txt_file):
        """Extract text from TXT file"""
        try:
            # Read as string
            text = str(txt_file.read(), "utf-8")
            return text.strip()
        except Exception as e:
            st.error(f"Error reading TXT file: {str(e)}")
            return ""
    
    def extract_text(self, uploaded_file):
        """Main method to extract text from any supported file type"""
        if uploaded_file is None:
            return ""
        
        file_extension = Path(uploaded_file.name).suffix.lower()
        
        with st.spinner(f"Extracting text from {uploaded_file.name}..."):
            if file_extension == '.pdf':
                return self.extract_text_from_pdf(uploaded_file)
            elif file_extension == '.docx':
                return self.extract_text_from_docx(uploaded_file)
            elif file_extension == '.doc':
                return self.extract_text_from_doc(uploaded_file)
            elif file_extension == '.txt':
                return self.extract_text_from_txt(uploaded_file)
            else:
                st.error(f"Unsupported file type: {file_extension}")
                return ""
    
    def save_uploaded_file(self, uploaded_file, save_directory="uploads"):
        """Save uploaded file to directory"""
        try:
            # Create directory if it doesn't exist
            os.makedirs(save_directory, exist_ok=True)
            
            file_path = os.path.join(save_directory, uploaded_file.name)
            
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            return file_path
        except Exception as e:
            st.error(f"Error saving file: {str(e)}")
            return None
    
    def get_file_info(self, uploaded_file):
        """Get information about uploaded file"""
        if uploaded_file is None:
            return {}
        
        file_extension = Path(uploaded_file.name).suffix.lower()
        file_size_mb = uploaded_file.size / (1024 * 1024)
        
        return {
            'name': uploaded_file.name,
            'size_mb': round(file_size_mb, 2),
            'type': uploaded_file.type,
            'extension': file_extension
        }

def create_file_uploader(label="Upload Resume", key=None):
    """Create a file uploader widget with validation"""
    uploaded_file = st.file_uploader(
        label,
        type=['pdf', 'docx', 'doc', 'txt'],
        help=f"Supported formats: {', '.join(config.ALLOWED_EXTENSIONS)}. Max size: {config.MAX_FILE_SIZE_MB}MB",
        key=key
    )
    
    if uploaded_file is not None:
        file_handler = FileHandler()
        is_valid, message = file_handler.validate_file(uploaded_file)
        
        if not is_valid:
            st.error(message)
            return None
        else:
            # Display file info
            file_info = file_handler.get_file_info(uploaded_file)
            st.success(f"✅ {file_info['name']} ({file_info['size_mb']} MB)")
    
    return uploaded_file